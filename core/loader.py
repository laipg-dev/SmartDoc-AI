"""
core/loader.py — Câu 1: Hỗ trợ PDF + DOCX
Thành viên B
"""
from pathlib import Path
from langchain_community.document_loaders import PDFPlumberLoader
from langchain_core.documents import Document


def load_pdf(file_path: str) -> list[Document]:
    """Load PDF bằng PDFPlumber — giữ layout và bảng biểu."""
    loader = PDFPlumberLoader(file_path)
    docs = loader.load()
    # Ghi đè source thành tên file (không phải full path)
    filename = Path(file_path).name
    for doc in docs:
        doc.metadata["source"] = filename
    return docs


def load_docx(file_path: str) -> list[Document]:
    """
    Câu 1: Load DOCX bằng python-docx.
    Trích xuất từng đoạn văn thành Document riêng.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("Cài python-docx: pip install python-docx")

    doc = DocxDocument(file_path)
    filename = Path(file_path).name
    documents = []

    # Gom các đoạn không rỗng thành chunks ~1000 ký tự
    buffer = []
    buffer_len = 0
    page_num = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        buffer.append(text)
        buffer_len += len(text)
        if buffer_len >= 800:
            documents.append(Document(
                page_content="\n".join(buffer),
                metadata={"source": filename, "page": page_num},
            ))
            buffer = []
            buffer_len = 0
            page_num += 1

    # Phần còn lại
    if buffer:
        documents.append(Document(
            page_content="\n".join(buffer),
            metadata={"source": filename, "page": page_num},
        ))

    return documents if documents else [
        Document(page_content="(Tài liệu trống)", metadata={"source": filename})
    ]


def load_file(file_path: str) -> list[Document]:
    """Tự động chọn loader theo đuôi file."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    else:
        raise ValueError(f"Định dạng không hỗ trợ: {ext}. Chỉ hỗ trợ .pdf và .docx")